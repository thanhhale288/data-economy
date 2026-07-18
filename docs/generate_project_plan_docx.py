"""Generate the complete Vietnamese project implementation plan as a DOCX."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Ke_hoach_thuc_hien_Du_an_Kinh_te_so_Nganh_Che_bien_Che_tao.docx"
DIAGRAM_DIR = ROOT / "diagram" / "output"

BLUE = "17365D"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=WHITE)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if row_index % 2 == 1:
                set_cell_shading(cells[i], LIGHT_GRAY)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_number(doc: Document, text: str, level: int = 0) -> None:
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(10)


def add_landscape_figure(doc: Document, filename: str, caption: str) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    path = DIAGRAM_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(9.7))
    add_caption(doc, caption)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color in [
        ("Title", 22, BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, BLUE),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    if "Figure Caption" not in styles:
        caption_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.name = "Times New Roman"
        caption_style.font.size = Pt(10)
        caption_style.font.italic = True

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Kế hoạch dự án Kinh tế số ngành Chế biến, Chế tạo — ")
        add_field(footer, "PAGE")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(24)
    run = p.add_run("KẾ HOẠCH THỰC HIỆN DỰ ÁN")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "NỀN TẢNG PHÂN TÍCH KINH TẾ SỐ\n"
        "NGÀNH CÔNG NGHIỆP CHẾ BIẾN, CHẾ TẠO VIỆT NAM"
    )
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)

    doc.add_paragraph()
    add_table(
        doc,
        ["Thuộc tính", "Nội dung"],
        [
            ["Lĩnh vực", "Khoa học dữ liệu và Kinh tế số"],
            ["Phạm vi ngành", "VSIC Section C — mã ngành 10–33"],
            ["Sản phẩm", "Web full-stack, cơ sở dữ liệu, pipeline tự động, ML/DL và benchmark"],
            ["Thời gian", "18 tuần (01 học kỳ)"],
            ["Phiên bản", "1.0"],
            ["Ngày lập", date.today().strftime("%d/%m/%Y")],
        ],
        widths=[4.5, 11.5],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Tài liệu tổng hợp từ docs/plan.md và bốn sơ đồ kiến trúc trong diagram/output."
    )
    run.italic = True
    run.font.size = Pt(10)
    doc.add_page_break()


def build_document() -> Document:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_heading("MỤC LỤC", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph(
        "Lưu ý: mở tài liệu bằng Microsoft Word và chọn Update Field để cập nhật mục lục."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(10)
    doc.add_page_break()

    doc.add_heading("DANH MỤC TỪ VIẾT TẮT", level=1)
    add_table(
        doc,
        ["Từ viết tắt", "Diễn giải"],
        [
            ["API", "Application Programming Interface"],
            ["BCTC", "Báo cáo tài chính"],
            ["CRISP-DM", "Cross-Industry Standard Process for Data Mining"],
            ["DL", "Deep Learning — học sâu"],
            ["GSO/NSO", "Cơ quan Thống kê Quốc gia Việt Nam"],
            ["IIP", "Index of Industrial Production — Chỉ số sản xuất công nghiệp"],
            ["ISIC", "International Standard Industrial Classification"],
            ["KTS", "Kinh tế số"],
            ["MAE/RMSE/MAPE", "Các thước đo sai số dự báo"],
            ["ML", "Machine Learning — học máy"],
            ["OECD", "Organisation for Economic Co-operation and Development"],
            ["SDMX", "Statistical Data and Metadata eXchange"],
            ["TMĐT", "Thương mại điện tử"],
            ["VA", "Value Added — giá trị gia tăng"],
            ["VDEI", "Vietnam Digital Economy Indicators (bộ chỉ tiêu đề xuất)"],
            ["VSIC", "Hệ thống ngành kinh tế Việt Nam"],
        ],
        widths=[4, 12],
    )

    doc.add_heading("1. TÓM TẮT ĐIỀU HÀNH", level=1)
    doc.add_paragraph(
        "Dự án xây dựng một nền tảng web hoàn chỉnh để thu thập, chuẩn hóa, phân tích "
        "và dự báo dữ liệu kinh tế số của ngành công nghiệp chế biến, chế tạo Việt Nam. "
        "Hệ thống kết hợp dữ liệu vĩ mô từ GSO/OECD với dữ liệu vi mô của doanh nghiệp "
        "niêm yết, báo cáo tài chính, website và kênh thương mại điện tử. Đầu ra gồm "
        "dashboard ngành, hồ sơ số hóa doanh nghiệp, ước lượng doanh thu trực tuyến và "
        "Digital VA dạng proxy, dự báo IIP, công cụ theo dõi pipeline và chức năng "
        "benchmark hiệu quả doanh nghiệp."
    )
    doc.add_paragraph(
        "Phạm vi MVP là khoảng 10 doanh nghiệp đại diện, triển khai trong 18 tuần. "
        "Dữ liệu crawl phải có nguồn gốc, thời điểm thu thập và trạng thái actual/"
        "estimate/forecast rõ ràng; hệ thống không được tự tạo số liệu khi nguồn thật "
        "không truy cập được. Dữ liệu fallback chỉ phục vụ kiểm thử kỹ thuật và phải "
        "được gắn nhãn riêng."
    )
    add_table(
        doc,
        ["Mục tiêu", "Kết quả đo lường"],
        [
            ["Tự động hóa dữ liệu", "Crawl theo lịch và thủ công; log trạng thái, số bản ghi, lỗi"],
            ["Đo hiện diện số", "Website, checkout, marketplace, social commerce và độ tin cậy match"],
            ["Ước lượng đóng góp KTS", "Online revenue proxy, Digital VA proxy, tỷ trọng theo ngành"],
            ["Dự báo xu hướng", "So sánh baseline, ML và DL bằng walk-forward validation"],
            ["Phục vụ người dùng", "5 module web và API có tài liệu OpenAPI"],
            ["Tái lập triển khai", "Docker Compose cho frontend, backend, worker, DB và Redis"],
        ],
        widths=[4.5, 11.5],
    )

    doc.add_heading("2. BỐI CẢNH VÀ ĐIỀU CHỈNH PROPOSAL", level=1)
    doc.add_paragraph(
        "Proposal ban đầu tập trung vào bán lẻ (VSIC Division 47) và dự báo tổng mức "
        "bán lẻ. Yêu cầu mới chuyển trọng tâm sang chế biến, chế tạo (VSIC Section C), "
        "đồng thời mở rộng từ nghiên cứu mô hình sang một sản phẩm phần mềm hoàn chỉnh."
    )
    add_table(
        doc,
        ["Nội dung cũ", "Điều chỉnh trong dự án"],
        [
            ["Bán lẻ và TMĐT", "Công nghiệp chế biến, chế tạo — VSIC Section C, mã 10–33"],
            ["Chỉ dữ liệu macro GSO/OECD", "Kết hợp macro và micro doanh nghiệp niêm yết"],
            ["Biến mục tiêu tổng mức bán lẻ", "IIP, giá trị gia tăng công nghiệp và chỉ tiêu số hóa DN"],
            ["Không định danh kênh bán số", "Phát hiện website, checkout, Shopee, TikTok/Lazada"],
            ["Superset/dashboard đơn lẻ", "Web React + FastAPI + PostgreSQL + pipeline worker"],
            ["Không có benchmark", "Module benchmark theo ý tưởng SingStat BITE"],
            ["Chỉ tiêu TMĐT bán lẻ", "Bộ 10 trụ cột VDEI Manufacturing"],
        ],
        widths=[6, 10],
    )

    doc.add_heading("3. MỤC TIÊU, CÂU HỎI NGHIÊN CỨU VÀ PHẠM VI", level=1)
    doc.add_heading("3.1. Mục tiêu tổng quát", level=2)
    doc.add_paragraph(
        "Thiết kế và triển khai nền tảng dữ liệu–AI có khả năng tự động thu thập, "
        "chuẩn hóa, lưu trữ, phân tích và trực quan hóa mức độ tham gia kinh tế số "
        "của ngành chế biến, chế tạo và các doanh nghiệp niêm yết đại diện."
    )
    doc.add_heading("3.2. Mục tiêu cụ thể", level=2)
    for item in [
        "Xây dựng pipeline thu thập định kỳ từ GSO, OECD, nguồn doanh nghiệp, BCTC và kênh số.",
        "Thiết lập ánh xạ ISIC–VSIC và mô hình dữ liệu có provenance.",
        "Xây dựng bộ chỉ tiêu số hóa và Digital VA proxy có công thức, giả định và độ tin cậy.",
        "Huấn luyện và so sánh mô hình thống kê, ML và DL cho dự báo IIP 3–6 tháng.",
        "Triển khai web app gồm Dashboard, Companies, Pipeline, ML Lab và Benchmark.",
        "Đóng gói toàn bộ hệ thống bằng Docker Compose và kiểm thử end-to-end.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("3.3. Câu hỏi nghiên cứu", level=2)
    for item in [
        "Dữ liệu GSO/OECD nào phản ánh phù hợp nhất quy mô và xu hướng Section C?",
        "Có thể đo hiện diện và mức tham gia TMĐT của doanh nghiệp niêm yết bằng dữ liệu công khai đến mức nào?",
        "Digital VA proxy thay đổi ra sao theo doanh thu online, biên lợi nhuận và mức độ áp dụng số?",
        "Các chỉ báo quốc tế có cải thiện dự báo IIP so với baseline chỉ dùng lịch sử IIP hay không?",
        "Doanh nghiệp đang ở percentile nào so với nhóm cùng ngành về hiệu quả tài chính và năng suất?",
    ]:
        add_number(doc, item)

    doc.add_heading("3.4. Phạm vi và ngoài phạm vi", level=2)
    add_table(
        doc,
        ["Trong phạm vi MVP", "Ngoài phạm vi MVP"],
        [
            ["VSIC Section C; khoảng 10 DN mẫu", "Toàn bộ doanh nghiệp Việt Nam"],
            ["Dữ liệu công khai và metadata nguồn", "Dữ liệu giao dịch nội bộ hoặc dữ liệu cá nhân"],
            ["Ước lượng doanh thu số dạng proxy", "Kiểm toán doanh thu TMĐT thực tế"],
            ["Forecast IIP 3–6 tháng", "Khẳng định quan hệ nhân quả kinh tế"],
            ["Web dashboard và benchmark prototype", "Hệ thống ra quyết định tín dụng/đầu tư"],
            ["PostgreSQL quan hệ", "Vector DB/RAG/LLM — chưa cần cho MVP"],
        ],
        widths=[8, 8],
    )

    doc.add_heading("4. ĐỐI TƯỢNG NGHIÊN CỨU VÀ QUY TẮC CHỌN DOANH NGHIỆP", level=1)
    doc.add_paragraph(
        "Đơn vị phân tích gồm cấp ngành (Section C và Division 10–33) và cấp doanh nghiệp "
        "niêm yết. Doanh nghiệp chỉ được đưa vào tập phân tích chính thức khi hoạt động "
        "chế biến, chế tạo là hoạt động chính hoặc có phân đoạn sản xuất được công bố đủ "
        "để tách số liệu. Mã VSIC phải được xác minh từ nguồn đăng ký hoặc báo cáo doanh nghiệp."
    )
    add_table(
        doc,
        ["Nhóm", "Mã minh họa", "Điều kiện sử dụng"],
        [
            ["Core manufacturing", "RAL, HPG, VNM, DGC", "Ưu tiên; xác minh VSIC và BCTC"],
            ["Tập đoàn đa ngành", "GVR, MSN, PNJ, REE", "Chỉ dùng phần sản xuất nếu tách được segment"],
            ["Cần rà soát phạm vi", "FPT, BWE", "Không mặc định là Section C; thay thế nếu không đạt tiêu chí"],
            ["Ứng viên thay thế", "BMP, NKG, DPM hoặc DN tương đương", "Chỉ chọn sau khi xác minh niêm yết và VSIC"],
        ],
        widths=[4, 4, 8],
    )
    doc.add_paragraph(
        "Lưu ý chất lượng phạm vi: danh sách seed trong MVP là danh sách phục vụ demo kỹ thuật, "
        "không tự động đồng nghĩa với mẫu nghiên cứu hợp lệ. Trước khi báo cáo kết quả, cần lập "
        "biên bản xác minh mã ngành và cơ cấu doanh thu cho từng doanh nghiệp."
    )

    doc.add_heading("5. KIẾN TRÚC TỔNG THỂ VÀ LUỒNG DỮ LIỆU", level=1)
    doc.add_paragraph(
        "Kiến trúc gồm năm lớp: nguồn dữ liệu, ingestion, clean/feature engineering, "
        "forecast/evaluation và store/serve. Các kiểm soát xuyên suốt gồm mapping VSIC–ISIC, "
        "provenance, quy tắc không tạo số liệu và nhật ký pipeline."
    )
    add_landscape_figure(
        doc,
        "01-data-pipeline.png",
        "Hình 1. Pipeline dữ liệu end-to-end của Manufacturing Data Economy Platform",
    )

    doc.add_heading("5.1. Các thành phần chính", level=2)
    add_table(
        doc,
        ["Lớp", "Thành phần", "Trách nhiệm"],
        [
            ["Sources", "GSO, OECD, HOSE/HNX, BCTC, website, marketplace", "Cung cấp dữ liệu macro và micro"],
            ["Ingestion", "GSO/OECD/company/marketplace crawlers", "Thu thập, parse, gắn metadata"],
            ["Transform", "Cleaning, metrics, feature engineering", "Chuẩn hóa, suy diễn proxy, tạo features"],
            ["ML", "Baseline, XGBoost, LSTM", "Train, đánh giá, dự báo"],
            ["Persistence", "PostgreSQL 16, model registry", "Lưu dữ liệu, jobs, predictions, metrics"],
            ["Serving", "FastAPI + React/Recharts", "API và giao diện người dùng"],
        ],
        widths=[3, 5, 8],
    )

    doc.add_heading("6. KẾ HOẠCH THU THẬP DỮ LIỆU", level=1)
    doc.add_heading("6.1. Luồng A — Dữ liệu macro GSO", level=2)
    add_table(
        doc,
        ["Dataset", "Phạm vi", "Tần suất", "Phương pháp", "Trường bắt buộc"],
        [
            ["IIP", "Section C và Division 10–33", "Tháng", "SDMX/XML; PX-Web fallback", "period, vsic, value, unit, source"],
            ["Shipment index", "Chế biến, chế tạo", "Tháng", "PX-Web/SDMX nếu có", "period, industry, index"],
            ["Inventory index", "Chế biến, chế tạo", "Tháng", "PX-Web/SDMX nếu có", "period, industry, index"],
            ["GDP/GRDP ngành", "Section C", "Quý/Năm", "PX-Web/Niên giám", "period, current/constant price, unit"],
            ["Thống kê DN", "Số DN, lao động, doanh thu", "Năm", "PX-Web/Niên giám", "year, industry, count/value"],
        ],
        widths=[3, 3.5, 2, 3.5, 5],
    )
    doc.add_heading("6.2. Luồng B — Dữ liệu doanh nghiệp", level=2)
    add_table(
        doc,
        ["Nhóm dữ liệu", "Các trường", "Nguồn", "Mức tự động"],
        [
            ["Định danh", "Mã CK, tên pháp lý, sàn, VSIC, website", "Sở GDCK, công bố DN", "Tự động + xác minh"],
            ["BCTC", "Doanh thu, PBT, tài sản, vốn CSH, chi phí, lao động", "PDF/HTML/XBRL", "Bán tự động"],
            ["Website", "URL, trạng thái, product/cart/checkout signals", "Website chính thức", "Tự động"],
            ["Marketplace", "Shop URL, listing, giá, sold count, rating", "Shopee/TikTok/Lazada", "Tự động có giới hạn"],
            ["Báo cáo thường niên", "Tỷ lệ online, đầu tư số, phân đoạn", "Investor relations", "PDF extraction + QA"],
        ],
        widths=[3, 6, 4, 3],
    )
    doc.add_heading("6.3. Luồng C — Dữ liệu OECD", level=2)
    add_table(
        doc,
        ["Chỉ tiêu", "Vai trò", "Tần suất", "Xử lý"],
        [
            ["Industrial Production Index", "Đối chiếu quốc tế / leading feature", "Tháng/Quý", "Mapping ISIC C và resample"],
            ["Business Confidence Index", "Chỉ báo dẫn dắt", "Tháng", "Lag 1–3 tháng/quý"],
            ["INDIGO", "Proxy độ mở thương mại số", "Theo kỳ công bố", "Gắn nhãn kỳ và nội suy có kiểm soát"],
            ["ICT investment by industry", "Proxy đầu tư số", "Năm/Quý", "Không nội suy nếu thiếu cơ sở"],
        ],
        widths=[4, 5, 3, 4],
    )

    doc.add_heading("6.4. Metadata và provenance bắt buộc", level=2)
    for item in [
        "source_name, source_url, publisher và dataset/table identifier;",
        "downloaded_at, observed_period và publication_date;",
        "unit, frequency, geographic scope, industry code và base year;",
        "extraction_method, parser_version và raw_file_hash;",
        "quality_status: raw/validated/rejected;",
        "value_status: actual/estimate/forecast/fallback;",
        "ghi chú định nghĩa, footnote, trang báo cáo và người QA.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("7. XỬ LÝ DỮ LIỆU VÀ BỘ CHỈ TIÊU", level=1)
    doc.add_heading("7.1. Quy trình làm sạch", level=2)
    add_number(doc, "Chuẩn hóa schema, kiểu dữ liệu, đơn vị và múi giờ.")
    add_number(doc, "Loại bản ghi trùng bằng source key và observation key.")
    add_number(doc, "Xử lý missing: gap ngắn có thể nội suy; gap dài phải gắn cờ và đánh giá thủ công.")
    add_number(doc, "Phát hiện outlier bằng IQR/Z-score nhưng không xóa tự động số liệu kinh tế hợp lệ.")
    add_number(doc, "Đồng bộ thời gian OECD–GSO; lưu cả dữ liệu gốc và dữ liệu resample.")
    add_number(doc, "Ánh xạ ISIC–VSIC ở cấp phù hợp; không ép mapping 1:1 nếu bản chất là many-to-many.")
    add_number(doc, "Entity resolution giữa tên DN, thương hiệu và shop; lưu confidence và bằng chứng.")

    doc.add_heading("7.2. Bộ chỉ tiêu VDEI Manufacturing", level=2)
    add_table(
        doc,
        ["Pillar", "Nội dung", "KPI cốt lõi"],
        [
            ["M1", "Quy mô & hiệu quả SXCN", "IIP, VA công nghiệp, tăng trưởng ngành"],
            ["M2", "Chuyển đổi số DN", "% có website, marketplace, ERP/IoT nếu có dữ liệu"],
            ["M3", "Doanh thu TMĐT", "Online revenue proxy / tổng doanh thu"],
            ["M4", "Kênh bán số", "Website riêng, marketplace, social commerce"],
            ["M5", "Hiệu quả số hóa", "Doanh thu/lao động, digital revenue/worker"],
            ["M6", "Đóng góp KTS", "Digital VA proxy và tỷ trọng"],
            ["M7", "Logistics số", "Kênh fulfillment, thời gian giao hàng nếu thu được"],
            ["M8", "Thanh toán số", "Tỷ lệ thanh toán số nếu công bố"],
            ["M9", "Xuất khẩu số", "Đơn hàng/kênh online quốc tế nếu công bố"],
            ["M10", "Cạnh tranh số", "Percentile benchmark cùng ngành"],
        ],
        widths=[2, 5, 9],
    )

    doc.add_heading("7.3. Công thức ước lượng", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Estimated Online Revenue = Σ(giá quan sát × số lượng đã bán quan sát)"
    )
    run.bold = True
    p = doc.add_paragraph()
    run = p.add_run(
        "Digital VA proxy = Online Revenue × Gross Margin "
        "+ Cost Saving proxy × Adoption Score − Digital Investment amortized"
    )
    run.bold = True
    doc.add_paragraph(
        "Đây là chỉ tiêu ước lượng, không phải giá trị gia tăng được kiểm toán. Mỗi kết quả "
        "phải đi kèm khoảng ước lượng hoặc confidence, phương pháp, kỳ dữ liệu và giả định. "
        "Không được cộng số lượng bán lũy kế của nhiều lần crawl như doanh số theo kỳ; cần "
        "lưu snapshot và lấy chênh lệch giữa hai thời điểm."
    )

    doc.add_heading("8. MÔ HÌNH DỮ LIỆU VÀ API", level=1)
    doc.add_heading("8.1. Các bảng nghiệp vụ", level=2)
    add_table(
        doc,
        ["Bảng", "Mục đích", "Khóa/quan hệ chính"],
        [
            ["vsic_codes", "Danh mục và mapping ISIC–VSIC", "vsic_code, parent_code"],
            ["companies", "Hồ sơ DN niêm yết", "stock_code, vsic_code"],
            ["financial_reports", "Chỉ tiêu BCTC theo kỳ", "company_id + period"],
            ["digital_presence", "Website/shop/social channels", "company_id + URL + crawled_at"],
            ["marketplace_listings", "Snapshot sản phẩm", "company_id + platform + product + time"],
            ["gso_macro", "Chuỗi GSO", "indicator + vsic + period"],
            ["oecd_indicators", "Chuỗi OECD", "indicator + country + period"],
            ["digital_metrics", "Online revenue/Digital VA proxy", "company_id + period"],
            ["model_registry", "Phiên bản và metric model", "model_name + version"],
            ["model_predictions", "Forecast và actual", "model + target + period"],
            ["pipeline_jobs", "Trạng thái vận hành", "job_name + run id"],
        ],
        widths=[4, 7, 5],
    )
    doc.add_heading("8.2. Nhóm API", level=2)
    add_table(
        doc,
        ["Nhóm", "Endpoint minh họa", "Chức năng"],
        [
            ["Dashboard", "GET /api/dashboard/*", "Summary, IIP, heatmap, OECD vs GSO"],
            ["Companies", "GET /api/companies/{code}", "Profile, channels, BCTC, metrics"],
            ["Macro", "GET /api/macro/gso|oecd", "Truy vấn chuỗi chỉ tiêu"],
            ["Pipeline", "POST /api/pipeline/trigger", "Khởi chạy crawl/transform/train"],
            ["Pipeline", "GET /api/pipeline/jobs", "Theo dõi trạng thái và lỗi"],
            ["ML", "POST /api/ml/train|forecast", "Huấn luyện và dự báo"],
            ["Benchmark", "POST /api/benchmark/compare", "Tính ratio và percentile"],
        ],
        widths=[3, 6, 7],
    )

    doc.add_heading("9. WORKFLOW VẬN HÀNH", level=1)
    doc.add_heading("9.1. Manual trigger", level=2)
    doc.add_paragraph(
        "Analyst chọn crawler trên giao diện Pipeline. Frontend gọi FastAPI để tạo "
        "PipelineJob, enqueue BackgroundTask, chạy các stage đã chọn, persist kết quả và "
        "cập nhật trạng thái để giao diện polling."
    )
    add_landscape_figure(
        doc,
        "02-manual-trigger-sequence.png",
        "Hình 2. Sequence khởi chạy pipeline thủ công",
    )

    doc.add_heading("9.2. Nightly pipeline", level=2)
    doc.add_paragraph(
        "Scheduler chạy lúc 02:00 hoặc nhận trigger thủ công. Nếu crawl thất bại, job được "
        "đánh dấu failed và dừng nhánh phụ thuộc. Nếu thành công, pipeline tính metrics, "
        "features, train/evaluate model, persist predictions và publish dữ liệu dashboard."
    )
    add_landscape_figure(
        doc,
        "03-nightly-pipeline-bpmn.png",
        "Hình 3. BPMN swimlane của pipeline chạy ban đêm",
    )

    doc.add_heading("10. AI/ML VÀ PHƯƠNG PHÁP ĐÁNH GIÁ", level=1)
    add_table(
        doc,
        ["Tầng", "Mô hình", "Target/Input", "Vai trò"],
        [
            ["Baseline", "Naive/seasonal naive, ARIMA/SARIMAX", "IIP history, seasonality", "Mốc so sánh có diễn giải"],
            ["ML", "XGBoost/LightGBM", "Lag, rolling, OECD, digital/financial features", "Quan hệ phi tuyến"],
            ["DL", "LSTM/GRU", "Chuỗi đa biến", "Phụ thuộc dài hạn; chỉ dùng nếu dữ liệu đủ"],
            ["Entity match", "Rule + RapidFuzz; classifier khi đủ nhãn", "Tên pháp lý, thương hiệu, shop", "Xác thực shop thuộc DN"],
            ["Anomaly", "IQR/Isolation Forest", "IIP/listing snapshots", "Cảnh báo, không tự xóa"],
        ],
        widths=[2.5, 4, 6, 3.5],
    )
    doc.add_heading("10.1. Thiết kế thí nghiệm", level=2)
    for item in [
        "Chia train/validation/test theo thời gian; tuyệt đối không random split chuỗi thời gian.",
        "Walk-forward validation với cùng horizon 1, 3 và 6 tháng.",
        "So sánh MAE, RMSE, MAPE; bổ sung sMAPE khi actual gần 0.",
        "Kiểm tra leakage: mọi feature tại tháng t chỉ dùng thông tin công bố trước hoặc tại t.",
        "Ablation test: IIP-only; IIP+OECD; IIP+OECD+digital.",
        "Ghi model version, data snapshot, code commit, hyperparameters và metric vào registry.",
        "LSTM chỉ được kết luận tốt hơn khi vượt baseline ổn định; không chọn model theo độ phức tạp.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("11. WEB APPLICATION VÀ BENCHMARK", level=1)
    add_table(
        doc,
        ["Module", "Chức năng chính", "Tiêu chí hoàn thành"],
        [
            ["Dashboard", "IIP, Digital VA, heatmap, OECD–GSO, forecast", "Filter kỳ/ngành; hiển thị source & updated_at"],
            ["Companies", "Profile, BCTC, website/shop, listings, metrics", "Case RAL đầy đủ và confidence"],
            ["Pipeline", "Trigger, job history, records, errors", "Không block request; trạng thái cập nhật"],
            ["ML Lab", "Model comparison, actual vs forecast, feature importance", "Cùng test window và metrics"],
            ["Benchmark", "Nhập BCTC, ratios, percentile", "ROA, ROE, current/equity ratio, worker metrics"],
        ],
        widths=[3, 8, 5],
    )
    doc.add_heading("11.1. Công thức benchmark", level=2)
    for item in [
        "ROA = Profit Before Interest and Tax / Average Total Assets;",
        "ROE = Profit Before Tax / Average Total Equity;",
        "Current Ratio = Current Assets / Current Liabilities;",
        "Equity Ratio = Total Equity / Total Assets;",
        "Revenue per Worker = Operating Revenue / Number of Employees;",
        "Profit per Worker = Profit Before Tax / Number of Employees.",
    ]:
        add_bullet(doc, item)
    doc.add_paragraph(
        "Percentile chỉ có ý nghĩa khi nhóm so sánh đủ số quan sát và đồng nhất về VSIC, "
        "kỳ báo cáo, quy mô và định nghĩa chỉ tiêu. Với mẫu 10 doanh nghiệp, kết quả benchmark "
        "chỉ là prototype, không được trình bày như chuẩn ngành đại diện."
    )

    doc.add_heading("12. TRIỂN KHAI VÀ HẠ TẦNG", level=1)
    doc.add_paragraph(
        "Hệ thống chạy trong Docker Compose. Frontend phục vụ tại cổng 5173, FastAPI tại "
        "8000, PostgreSQL tại 5432 và Redis tại 6379. Worker thực hiện scheduler và các "
        "pipeline job; volume PostgreSQL đảm bảo dữ liệu tồn tại sau khi container restart."
    )
    add_landscape_figure(
        doc,
        "04-docker-compose-network.png",
        "Hình 4. Mạng triển khai Docker Compose",
    )
    doc.add_heading("12.1. Môi trường", level=2)
    add_table(
        doc,
        ["Môi trường", "Mục đích", "Yêu cầu"],
        [
            ["Local", "Phát triển nhanh", "SQLite có thể dùng tạm; dữ liệu demo gắn nhãn"],
            ["Integration", "Kiểm thử end-to-end", "PostgreSQL + Redis + worker"],
            ["Demo", "Trình bày học kỳ", "Docker Compose, seed kiểm soát, backup DB"],
            ["Production future", "Mở rộng vận hành", "Secrets, HTTPS, monitoring, managed DB"],
        ],
        widths=[3, 6, 7],
    )
    doc.add_heading("12.2. Bảo mật và quản trị", level=2)
    for item in [
        "Không hard-code mật khẩu; dùng .env/secrets và thay credential mặc định.",
        "Chỉ crawl dữ liệu công khai; tuân thủ robots.txt, điều khoản nền tảng và rate limit.",
        "Không thu thập dữ liệu cá nhân không cần thiết.",
        "Sanitize URL/input; giới hạn crawler allowlist để tránh SSRF.",
        "CORS theo domain demo; thêm authentication cho trigger/train trong môi trường công khai.",
        "Backup PostgreSQL trước demo và lưu raw artifact bất biến.",
        "Ghi audit log cho trigger thủ công và thay đổi dữ liệu đã QA.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("13. LỘ TRÌNH 18 TUẦN", level=1)
    add_table(
        doc,
        ["Giai đoạn", "Tuần", "Công việc", "Đầu ra/Gate"],
        [
            ["1. Nền tảng & macro", "1–5", "Scaffold; DB; VSIC mapping; GSO/OECD crawlers; seed", "API health; macro data có provenance"],
            ["2. DN & hiện diện số", "6–10", "Xác minh mẫu; BCTC; website; marketplace; matcher", "10 profile hoặc biên bản loại/thay thế"],
            ["3. Clean & ML", "11–14", "Cleaning; feature store; baseline/XGB/LSTM; registry", "Bảng metric walk-forward; không leakage"],
            ["4. Web & tích hợp", "15–17", "5 modules; API integration; E2E; performance", "Demo end-to-end và test report"],
            ["5. Benchmark & báo cáo", "18", "Benchmark; QA số liệu; docs; presentation", "Bộ deliverable cuối kỳ"],
        ],
        widths=[3, 2, 7, 4],
    )
    doc.add_heading("13.1. Kế hoạch tuần chi tiết", level=2)
    add_table(
        doc,
        ["Tuần", "Nhiệm vụ trọng tâm", "Tiêu chí ra khỏi tuần"],
        [
            ["1", "Chốt scope, data dictionary, tiêu chí mẫu", "Scope & source inventory được duyệt"],
            ["2", "Docker, DB schema, migrations, API skeleton", "Stack chạy local/integration"],
            ["3", "GSO IIP crawler + raw persistence", "Có raw artifact và validated table"],
            ["4", "OECD SDMX + mapping thời gian/ngành", "Chuỗi OECD hợp nhất được"],
            ["5", "Dashboard macro sơ bộ + pipeline jobs", "Macro vertical slice hoàn chỉnh"],
            ["6", "Xác minh 10 DN và VSIC", "Danh sách hợp lệ hoặc replacement"],
            ["7", "BCTC parser + manual QA workflow", "Chỉ tiêu BCTC mẫu đối chiếu được"],
            ["8", "Website detector", "Signals và evidence lưu DB"],
            ["9", "Marketplace snapshots + rate limit", "Listing snapshots có timestamp"],
            ["10", "Shop matcher + Digital metrics v1", "Confidence và công thức proxy"],
            ["11", "Cleaning + provenance QA", "Data quality report"],
            ["12", "Feature engineering + leakage tests", "Feature dataset versioned"],
            ["13", "Baseline + XGBoost", "Walk-forward metrics"],
            ["14", "LSTM + model comparison", "Model recommendation có căn cứ"],
            ["15", "Dashboard + Companies", "Hai module E2E"],
            ["16", "Pipeline + ML Lab", "Trigger/status/forecast chạy được"],
            ["17", "Benchmark + system testing", "Acceptance tests pass"],
            ["18", "Freeze data/model; report; demo rehearsal", "Release candidate và backup"],
        ],
        widths=[1.5, 9, 5.5],
    )

    doc.add_heading("14. KIỂM THỬ VÀ TIÊU CHÍ NGHIỆM THU", level=1)
    add_table(
        doc,
        ["Hạng mục", "Kiểm thử", "Tiêu chí nghiệm thu"],
        [
            ["Crawler", "Unit parser + integration source", "Không crash; retry; provenance đầy đủ; không tạo số giả"],
            ["Data", "Schema, duplicate, range, freshness", "Critical checks pass; lỗi được quarantine"],
            ["Entity match", "Gold labels thủ công", "Precision mục tiêu ≥90% trên tập đánh giá nhỏ"],
            ["ML", "Walk-forward + leakage", "Có baseline; metric tái lập; model card"],
            ["API", "Contract, error, pagination", "2xx/4xx đúng; OpenAPI đầy đủ"],
            ["Frontend", "Navigation, empty/loading/error", "5 module hoạt động; responsive cơ bản"],
            ["Pipeline", "Manual + nightly", "Job status, records, error và rerun hoạt động"],
            ["Benchmark", "Formula unit tests", "Kết quả khớp tính tay; cảnh báo mẫu nhỏ"],
            ["Deployment", "Fresh compose up", "Khởi động từ máy sạch theo README"],
        ],
        widths=[3, 6, 7],
    )
    doc.add_heading("14.1. Definition of Done", level=2)
    for item in [
        "Code đã review, test liên quan pass và không còn lỗi nghiêm trọng.",
        "Nguồn dữ liệu, công thức và giả định được ghi trong tài liệu.",
        "Có screenshot/demo và endpoint tương ứng cho chức năng.",
        "Không dùng fallback hoặc estimate mà thiếu nhãn.",
        "Có hướng dẫn chạy lại từ môi trường sạch.",
        "Có backup DB, model artifact và data snapshot dùng cho demo.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("15. RỦI RO VÀ BIỆN PHÁP GIẢM THIỂU", level=1)
    add_table(
        doc,
        ["Rủi ro", "Xác suất", "Tác động", "Giảm thiểu"],
        [
            ["GSO thay đổi endpoint/UI", "Trung bình", "Cao", "Ưu tiên SDMX; adapter; lưu raw; test parser"],
            ["Marketplace chặn crawl", "Cao", "Cao", "Rate limit; cache; sample nhỏ; nguồn công khai thay thế"],
            ["Vi phạm điều khoản nền tảng", "Trung bình", "Cao", "Rà soát ToS; không bypass; dừng crawler khi không được phép"],
            ["Sai match shop–DN", "Trung bình", "Cao", "Confidence, evidence, manual QA, gold set"],
            ["BCTC PDF khó parse", "Cao", "Trung bình", "Ưu tiên HTML/XBRL; extraction template; manual verification"],
            ["Mẫu DN không thuộc Section C", "Trung bình", "Cao", "Eligibility gate và danh sách thay thế"],
            ["Thiếu doanh thu TMĐT", "Cao", "Cao", "Proxy có khoảng; không khẳng định actual"],
            ["Dữ liệu chuỗi ngắn cho DL", "Cao", "Trung bình", "Baseline-first; giảm model; không overclaim"],
            ["Data leakage", "Trung bình", "Cao", "Publication-time features và walk-forward"],
            ["Docker/resource không ổn định", "Trung bình", "Trung bình", "Resource limits; train offline; lưu artifact"],
            ["Lộ secrets/API", "Thấp", "Cao", ".env, rotate credentials, không commit secrets"],
        ],
        widths=[5, 2, 2, 7],
    )

    doc.add_heading("16. DELIVERABLE CUỐI HỌC KỲ", level=1)
    deliverables = [
        "Mã nguồn backend, frontend, crawlers, pipeline và ML/DL.",
        "Cơ sở dữ liệu PostgreSQL với migration, mapping và dữ liệu đã gắn provenance.",
        "Pipeline manual và nightly, có nhật ký trạng thái.",
        "Web app 5 module: Dashboard, Companies, Pipeline, ML Lab, Benchmark.",
        "Ba nhóm mô hình cùng báo cáo đánh giá trên cùng test window.",
        "Case study Rạng Đông: VSIC, BCTC, website/shop, online revenue và Digital VA proxy.",
        "Bộ sơ đồ draw.io/PNG: data pipeline, manual sequence, nightly BPMN và network.",
        "Báo cáo nghiên cứu, data dictionary, model card, test report và README triển khai.",
        "Slide demo, data/model snapshot và bản backup dùng cho buổi bảo vệ.",
    ]
    for item in deliverables:
        add_number(doc, item)

    doc.add_heading("17. KẾ HOẠCH DEMO", level=1)
    add_number(doc, "Mở Dashboard và giải thích IIP, nguồn dữ liệu, thời gian cập nhật.")
    add_number(doc, "Mở hồ sơ RAL, chỉ ra website/shop, listing snapshot và confidence.")
    add_number(doc, "Trigger một crawler từ Pipeline Monitor và theo dõi trạng thái.")
    add_number(doc, "Mở ML Lab, so sánh baseline/XGBoost/LSTM trên cùng test window.")
    add_number(doc, "Chạy forecast 3–6 tháng và giải thích uncertainty.")
    add_number(doc, "Nạp BCTC mẫu vào Benchmark và xem ratio/percentile.")
    add_number(doc, "Kết thúc bằng limitation: proxy, mẫu nhỏ, marketplace access và hướng mở rộng.")

    doc.add_heading("18. HƯỚNG PHÁT TRIỂN SAU MVP", level=1)
    for item in [
        "Mở rộng sang toàn bộ doanh nghiệp chế biến, chế tạo niêm yết đủ điều kiện.",
        "Bổ sung dữ liệu hải quan/xuất khẩu, logistics, thanh toán số và điều tra doanh nghiệp.",
        "Dùng probabilistic forecasting và prediction interval.",
        "Xây benchmark theo size band, VSIC và năm với mẫu đủ lớn.",
        "Thêm data catalog, lineage, observability và orchestration chuyên dụng.",
        "Chỉ cân nhắc Vector DB/RAG khi có nhu cầu hỏi đáp trên BCTC, báo cáo GSO/OECD và tài liệu; không phải thành phần bắt buộc của pipeline hiện tại.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("19. TÀI LIỆU VÀ TỆP THAM CHIẾU", level=1)
    add_table(
        doc,
        ["Tệp/Nguồn", "Vai trò"],
        [
            ["docs/plan.md", "Kế hoạch gốc và phạm vi 18 tuần"],
            ["docs/proposal-v2.md", "Phương pháp nghiên cứu và kiến trúc cập nhật"],
            ["diagram/output/01-data-pipeline.png", "Pipeline end-to-end"],
            ["diagram/output/02-manual-trigger-sequence.png", "Sequence manual trigger"],
            ["diagram/output/03-nightly-pipeline-bpmn.png", "BPMN/swimlane scheduled pipeline"],
            ["diagram/output/04-docker-compose-network.png", "Kiến trúc triển khai"],
            ["README.md", "Hướng dẫn chạy project"],
            ["SingStat BITE", "Tham chiếu ý tưởng benchmark hiệu quả DN"],
            ["GSO/NSO và OECD SDMX", "Nguồn dữ liệu thống kê chính thức"],
        ],
        widths=[7, 9],
    )

    doc.add_heading("PHỤ LỤC A — MA TRẬN TRUY VẾT YÊU CẦU", level=1)
    add_table(
        doc,
        ["Yêu cầu", "Thành phần", "Bằng chứng nghiệm thu"],
        [
            ["Web hoàn chỉnh", "React + FastAPI", "5 module và OpenAPI"],
            ["Database", "PostgreSQL + SQLAlchemy", "Schema, migration, data dictionary"],
            ["Data pipeline", "Crawlers + worker", "Manual/nightly job logs"],
            ["AI/ML/DL", "Baseline + XGBoost + LSTM", "Model registry và metric report"],
            ["Crawl ngành C", "GSO/OECD + VSIC mapping", "Raw/validated observations"],
            ["DN niêm yết", "Companies/BCTC/digital presence", "Danh sách eligibility và profile"],
            ["Doanh thu số", "Marketplace snapshots", "Proxy formula + provenance + confidence"],
            ["Benchmark", "Benchmark API/UI", "Ratios, percentile và cảnh báo mẫu"],
        ],
        widths=[4, 6, 6],
    )

    return doc


def main() -> None:
    doc = build_document()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
